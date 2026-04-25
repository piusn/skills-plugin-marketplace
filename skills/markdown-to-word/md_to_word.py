"""
Markdown to Word Document Converter
====================================
Converts a Markdown file to a professionally formatted Word (.docx) document.

Usage:
    python md_to_word.py <input.md> <output.docx>

Requires: pip install python-docx
"""

import sys
import re
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Customizable constants ──────────────────────────────────────────────
HEADER_BG = '0078D4'
ALT_ROW_BG = 'F0F6FC'
HEADING_COLOR = RGBColor(0x00, 0x78, 0xD4)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
HINT_COLOR = RGBColor(0x66, 0x66, 0x66)
CALLOUT_COLOR = RGBColor(0xD8, 0x3B, 0x01)
FONT_BODY = 'Segoe UI'
FONT_HEADING = 'Segoe UI Semibold'
FONT_CODE = 'Consolas'
BODY_SIZE = Pt(10.5)
HEADING_SIZES = {1: Pt(22), 2: Pt(16), 3: Pt(13)}
HEADING_SPACE_BEFORE = {1: Pt(24), 2: Pt(18), 3: Pt(14)}
HEADING_SPACE_AFTER = {1: Pt(12), 2: Pt(8), 3: Pt(6)}
PAGE_MARGIN = Cm(2.54)


def setup_styles(doc):
    """Configure document-wide styles."""
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = BODY_SIZE
    style.font.color.rgb = BODY_COLOR
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = FONT_HEADING
        hs.font.color.rgb = HEADING_COLOR
        hs.font.size = HEADING_SIZES[level]
        hs.paragraph_format.space_before = HEADING_SPACE_BEFORE[level]
        hs.paragraph_format.space_after = HEADING_SPACE_AFTER[level]

    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN


def add_styled_table(doc, headers, rows):
    """Add a table with blue header row and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = FONT_HEADING
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text).strip())
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.font.color.rgb = BODY_COLOR
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ALT_ROW_BG}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph('')  # spacing
    return table


def add_divider(doc):
    """Add a horizontal line divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def apply_inline_formatting(paragraph, text):
    """Parse inline Markdown formatting (bold, italic, code) and add runs."""
    # Pattern to match **bold**, *italic*, _italic_, `code`, and plain text
    pattern = r'(\*\*(.+?)\*\*|_\((.+?)\)_|_(.+?)_|\*(.+?)\*|`(.+?)`)'
    last_end = 0

    for match in re.finditer(pattern, text):
        # Add plain text before this match
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            if plain:
                run = paragraph.add_run(plain)
                run.font.size = BODY_SIZE
                run.font.name = FONT_BODY

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.size = BODY_SIZE
            run.font.name = FONT_HEADING
        elif match.group(3):  # _(italic with parens)_
            run = paragraph.add_run(f'({match.group(3)})')
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.font.color.rgb = HINT_COLOR
        elif match.group(4):  # _italic_
            run = paragraph.add_run(match.group(4))
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.font.color.rgb = HINT_COLOR
        elif match.group(5):  # *italic*
            run = paragraph.add_run(match.group(5))
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.font.color.rgb = HINT_COLOR
        elif match.group(6):  # `code`
            run = paragraph.add_run(match.group(6))
            run.font.name = FONT_CODE
            run.font.size = Pt(10)

        last_end = match.end()

    # Remaining plain text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            run = paragraph.add_run(remaining)
            run.font.size = BODY_SIZE
            run.font.name = FONT_BODY


def parse_table_row(line):
    """Parse a Markdown table row into cells."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]


def is_separator_row(line):
    """Check if a table line is a separator row (e.g., |---|---|)."""
    return bool(re.match(r'^[\s|:-]+$', line))


def convert_md_to_docx(input_path, output_path):
    """Main conversion function."""
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            add_divider(doc)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```

            p = doc.add_paragraph()
            code_text = '\n'.join(code_lines)
            run = p.add_run(code_text)
            run.font.name = FONT_CODE
            run.font.size = Pt(9)
            run.font.color.rgb = BODY_COLOR
            # Light gray background via shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            p._p.get_or_add_pPr().append(shading)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            continue

        # Table detection
        if '|' in stripped and not stripped.startswith('>'):
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                headers = parse_table_row(table_lines[0])
                data_rows = []
                for tl in table_lines[1:]:
                    if not is_separator_row(tl):
                        row = parse_table_row(tl)
                        # Pad or trim to match header count
                        while len(row) < len(headers):
                            row.append('')
                        row = row[:len(headers)]
                        data_rows.append(row)
                if data_rows:
                    add_styled_table(doc, headers, data_rows)
                elif headers:
                    add_styled_table(doc, headers, [[''] * len(headers)])
            continue

        # Checklist items
        checklist_match = re.match(r'^-\s*\[[ x]\]\s*(.+)$', stripped)
        if checklist_match:
            p = doc.add_paragraph()
            run = p.add_run('☐  ' + checklist_match.group(1))
            run.font.size = BODY_SIZE
            p.paragraph_format.left_indent = Cm(1.0)
            i += 1
            continue

        # Bullet with bold prefix: - **Label:** text
        bold_bullet_match = re.match(r'^[-*]\s+\*\*(.+?)\*\*\s*(.*)', stripped)
        if bold_bullet_match:
            label = bold_bullet_match.group(1)
            rest = bold_bullet_match.group(2)
            p = doc.add_paragraph()
            run_b = p.add_run(label + ' ')
            run_b.bold = True
            run_b.font.size = BODY_SIZE
            run_b.font.name = FONT_HEADING
            if rest:
                apply_inline_formatting(p, rest)
            i += 1
            continue

        # Regular bullet
        bullet_match = re.match(r'^[-*]\s+(.+)$', stripped)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            apply_inline_formatting(p, bullet_match.group(1))
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            apply_inline_formatting(p, text)
            # Apply callout color to all runs
            for run in p.runs:
                if not run.bold:
                    run.font.color.rgb = CALLOUT_COLOR
            i += 1
            continue

        # Italic instruction line
        italic_match = re.match(r'^_(.+)_$', stripped)
        if italic_match:
            p = doc.add_paragraph()
            run = p.add_run(italic_match.group(1))
            run.italic = True
            run.font.color.rgb = HINT_COLOR
            run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        apply_inline_formatting(p, stripped)
        i += 1

    # Ensure output directory exists
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc.save(output_path)
    print(f'Converted: {input_path} -> {output_path}')


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python md_to_word.py <input.md> <output.docx>')
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    if not os.path.exists(input_file):
        print(f'Error: Input file not found: {input_file}')
        sys.exit(1)

    convert_md_to_docx(input_file, output_file)
